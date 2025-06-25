
import os
import io
import logging
from typing import List, Tuple

import numpy as np
import streamlit as st
from sentence_transformers import SentenceTransformer
import faiss
from PyPDF2 import PdfReader
import docx
from dotenv import load_dotenv

# Page configuration (must be first Streamlit command)
st.set_page_config(page_title="🔍 AI-Powered Doc Search", layout="wide")

# Load environment variables
load_dotenv()

# Default configuration from environment
default_docs_dir = os.getenv("DOCS_DIR", "docs")
default_chunk_size = int(os.getenv("CHUNK_SIZE", 500))
default_overlap = int(os.getenv("CHUNK_OVERLAP", 50))
default_model_name = os.getenv("EMBEDDING_MODEL", "all-MiniLM-L6-v2")
default_top_k = int(os.getenv("TOP_K", 5))

# Logging setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Sidebar: Data Source Selection
st.sidebar.header("📂 Data Source")
# Session state storage for directory
if "docs_dir" not in st.session_state:
    st.session_state.docs_dir = default_docs_dir

def _browse_directory():
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        folder = filedialog.askdirectory()
        if folder:
            st.session_state.docs_dir = folder
    except Exception as e:
        st.error(f"Directory dialog failed: {e}")

st.sidebar.button("Browse Local Directory", on_click=_browse_directory)
# Text input fallback
docs_dir = st.sidebar.text_input("Or type local docs directory", st.session_state.docs_dir)
# File uploader fallback
uploaded_files = st.sidebar.file_uploader(
    "Or upload PDF/DOCX files", type=["pdf", "docx"], accept_multiple_files=True
)

# Sidebar: Indexing Settings
st.sidebar.header("⚙️ Indexing Settings")
chunk_size = st.sidebar.number_input(
    "Chunk size (chars)", min_value=100, value=default_chunk_size
)
overlap = st.sidebar.number_input(
    "Chunk overlap (chars)", min_value=0, value=default_overlap
)
model_name = st.sidebar.text_input(
    "Embedding model", value=default_model_name
)
top_k = st.sidebar.number_input(
    "Top K results", min_value=1, max_value=20, value=default_top_k
)

# Load embedding model (cached)
@st.cache_resource
def load_embedding_model(name: str) -> SentenceTransformer:
    logger.info(f"Loading embedding model: {name}")
    return SentenceTransformer(name)

# Build FAISS index from a directory (cached)
@st.cache_resource(show_spinner=False)
def build_index_from_dir(
    _model: SentenceTransformer,
    path: str,
    chunk_size: int,
    overlap: int,
) -> Tuple[faiss.IndexFlatIP, List[Tuple[str, str]]]:
    texts, meta = [], []
    for root, _, files in os.walk(path):
        for fname in files:
            if not fname.lower().endswith((".pdf", ".docx")):
                continue
            full = os.path.join(root, fname)
            try:
                if fname.lower().endswith(".pdf"):
                    reader = PdfReader(full)
                    raw = "".join(p.extract_text() or "" for p in reader.pages)
                else:
                    doc = docx.Document(full)
                    raw = "".join(p.text for p in doc.paragraphs)
            except Exception as e:
                logger.warning(f"Failed reading {full}: {e}")
                continue
            start = 0
            while start < len(raw):
                chunk = raw[start : start + chunk_size]
                texts.append(chunk)
                meta.append((fname, chunk))
                start += chunk_size - overlap
    dim = _model.get_sentence_embedding_dimension()
    if not texts:
        st.error(f"No documents found in '{path}' and no files uploaded.")
        st.stop()
    embeddings = _model.encode(texts, convert_to_numpy=True, show_progress_bar=True)
    faiss.normalize_L2(embeddings)
    index = faiss.IndexFlatIP(dim)
    index.add(embeddings)
    return index, meta

# Build FAISS index from uploaded files
def build_index_from_uploads(
    _model,
    files,
    chunk_size: int,
    overlap: int,
) -> Tuple[faiss.IndexFlatIP, List[Tuple[str, str]]]:
    texts, meta = [], []
    for file in files:
        name = file.name
        data = file.read()
        try:
            if name.lower().endswith(".pdf"):
                reader = PdfReader(io.BytesIO(data))
                raw = "".join(p.extract_text() or "" for p in reader.pages)
            else:
                doc = docx.Document(io.BytesIO(data))
                raw = "".join(p.text for p in doc.paragraphs)
        except Exception as e:
            logger.warning(f"Failed reading {name}: {e}")
            continue
        start = 0
        while start < len(raw):
            chunk = raw[start : start + chunk_size]
            texts.append(chunk)
            meta.append((name, chunk))
            start += chunk_size - overlap
    dim = _model.get_sentence_embedding_dimension()
    if not texts:
        st.error("No valid uploads to index.")
        st.stop()
    embeddings = _model.encode(texts, convert_to_numpy=True, show_progress_bar=True)
    faiss.normalize_L2(embeddings)
    index = faiss.IndexFlatIP(dim)
    index.add(embeddings)
    return index, meta

# Initialize embedding model and FAISS index
model = load_embedding_model(model_name)
if uploaded_files:
    index, metadata = build_index_from_uploads(model, uploaded_files, chunk_size, overlap)
elif os.path.isdir(docs_dir):
    index, metadata = build_index_from_dir(model, docs_dir, chunk_size, overlap)
else:
    st.error(f"Directory '{docs_dir}' not found. Please correct the path or upload files.")
    st.stop()

# Main UI
st.title("🔍 AI-Powered Document Search")

st.markdown(f"""**Directory**: `{docs_dir}`
**Chunk**: {chunk_size} chars | **Overlap**: {overlap} chars
**Model**: {model_name} | **Top K**: {top_k}""")

query = st.text_input("Enter your search query:")
if query:
    q_emb = model.encode(query, convert_to_numpy=True)
    faiss.normalize_L2(q_emb.reshape(1, -1))
    dists, idxs = index.search(q_emb.reshape(1, -1), top_k)
    st.subheader(f"Top {top_k} Results for: '{query}'")
    for rank, (dist, idx) in enumerate(zip(dists[0], idxs[0]), start=1):
        fname, snippet = metadata[idx]
        st.markdown(f"**{rank}. {fname}** (score: {dist:.3f})")
        st.write(snippet)
        st.markdown("---")

